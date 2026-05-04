"""
extraer_sgi.py
--------------
Extrae los campos del formato SGI "Seguimiento de la Gestión del Curso"
desde archivos .docx sueltos, dentro de .zip, o en carpetas completas.
Al finalizar genera un archivo Excel con todos los datos.

Uso:
    python extraer_sgi.py                  <- usa la carpeta '../data' por defecto
    python extraer_sgi.py C:/ruta/carpeta  <- carpeta personalizada
    python extraer_sgi.py C:/archivo.zip   <- un solo ZIP
    python extraer_sgi.py C:/archivo.docx  <- un solo DOCX
"""

import os
import re
import sys
import zipfile
import tempfile
from pathlib import Path

import win32com.client
from docx import Document
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment


# ---------------------------------------------------------------------------
# Configuración
# ---------------------------------------------------------------------------

RUTA_POR_DEFECTO = r"..\data"

ENCABEZADO_EXCEL = [
    "Archivo", "Revisión", "Docente", "Asignatura",
    "Cantidad De Estudiantes Inscritos",
    "Periodo", "Unidades", "Captura de Calificaciones",
    "Aprobados", "Reprobados", "Evidencias de Aprendizaje",
    "% De Avance", "Producto No Conforme",
]


# ---------------------------------------------------------------------------
# Helpers de texto
# ---------------------------------------------------------------------------

def limpiar(texto: str) -> str:
    """Quita guiones de relleno, saltos de línea y espacios múltiples."""
    texto = texto.replace("\n", " ").replace("\r", " ").strip()
    texto = re.sub(r"_{2,}", "", texto)
    texto = re.sub(r"\s{2,}", " ", texto)
    return texto.strip()


def extraer_campo(texto: str, patron: str) -> str:
    """Busca 'patron: VALOR' tolerando variaciones de mayúsculas y espacios."""
    match = re.search(patron + r"[:\s]+(.+?)(?=\s{2,}|$)", texto, re.IGNORECASE)
    return limpiar(match.group(1)) if match else ""


def es_placeholder(valor: str) -> bool:
    """Retorna True si el valor es un marcador sin llenar del instructivo: (2), (3), etc."""
    return bool(re.fullmatch(r"\(\d{1,2}\)", valor.strip()))


# ---------------------------------------------------------------------------
# Extracción del encabezado (Revisión, Docente, Asignatura, Estudiantes)
# ---------------------------------------------------------------------------

def extraer_encabezado(doc: Document) -> dict:
    datos = {
        "Revisión":                          "",
        "Docente":                           "",
        "Asignatura":                        "",
        "Cantidad De Estudiantes Inscritos": "",
    }

    # Revisión está en la tabla del header de sección del documento
    for section in doc.sections:
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    match = re.search(r"Revisi[oó]n[:\s]*(\d+)", cell.text, re.IGNORECASE)
                    if match:
                        datos["Revisión"] = match.group(1)

    # Docente, Asignatura y Estudiantes en párrafos de texto libre
    # Flag para el caso donde DOCENTE: está vacío y el nombre va en el siguiente párrafo
    esperando_nombre_docente = False

    for parrafo in doc.paragraphs:
        texto = parrafo.text
        if not texto.strip():
            continue

        # Si el párrafo anterior tenía DOCENTE: vacío, este párrafo es el nombre
        if esperando_nombre_docente and not re.search(r"ASIGNATURA|DEPARTAMENTO|CANTIDAD", texto, re.IGNORECASE):
            val = re.sub(r"\s{2,}", " ", texto.strip().strip("_").strip())
            if not es_placeholder(val):
                datos["Docente"] = val
            esperando_nombre_docente = False
            continue

        if not datos["Docente"]:
            # Salta guiones bajos iniciales después de ':', acepta espacio antes de ':'
            m = re.search(r"DOCENTE\s*:\s*_*\s*(?P<name>[^_\n\r\t]+?)(?=_{2,}|\s*$)", texto, re.IGNORECASE)
            if re.search(r"DOCENTE\s*:", texto, re.IGNORECASE):
                if m and m.group("name").strip():
                    val = re.sub(r"\s{2,}", " ", m.group("name").strip())
                    if not es_placeholder(val):
                        datos["Docente"] = val
                else:
                    esperando_nombre_docente = True  # nombre en el siguiente párrafo

        if not datos["Asignatura"]:
            # Captura todo el texto entre 'ASIGNATURA' y 'CANTIDAD'
            m = re.search(r"ASIGNATURA(?:\s+Y\s+GRUPO)?\s*:\s*(?P<name>.+?)(?=CANTIDAD|\s{5,}|$)", texto, re.IGNORECASE)
            if m:
                # Limpiamos guiones bajos, comillas raras y espacios extra
                val = m.group("name").replace("_", " ").strip()
                val = re.sub(r"\s{2,}", " ", val)  # Normaliza espacios
                if not es_placeholder(val):
                    datos["Asignatura"] = val

        if not datos["Cantidad De Estudiantes Inscritos"]:
            # Captura solo dígitos, ignora guiones bajos al final
            m = re.search(r"CANTIDAD\s+DE\s+ESTUDIANTES?\s+INSCRITOS?[:\s]+(?:\(\d+\))?[\s_]*(?P<num>\d+)", texto, re.IGNORECASE)
            if m:
                val = m.group(1)
                if not es_placeholder(val):
                    datos["Cantidad De Estudiantes Inscritos"] = val

    return datos


# ---------------------------------------------------------------------------
# Extracción de la tabla de seguimiento
# ---------------------------------------------------------------------------

_KW_TABLA = {"UNIDADES", "APROBADOS", "REPROBADOS", "AVANCE", "CONFORME"}


def es_tabla_seguimiento(table) -> bool:
    """Identifica la tabla de datos por palabras clave en sus encabezados."""
    encabezado = " ".join(cell.text.upper() for cell in table.rows[0].cells)
    return sum(1 for kw in _KW_TABLA if kw in encabezado) >= 3


def fila_es_dato(celdas: list) -> bool:
    """
    Acepta filas de datos reales y excluye las filas de encabezados de la tabla.
    """
    if not celdas or len(celdas) < 2:
        return False

    # Limpiamos el texto de la celda de la unidad
    unidad = celdas[1].strip().upper()

    # --- FILTRO DE EXCLUSIÓN ---
    # Si la celda contiene exactamente estas palabras, es un encabezado, NO un dato.
    palabras_prohibidas = {"UNIDADES", "PERIODO", "APROBADOS", "REPROBADOS", "CAPACITACIÓN"}
    if unidad in palabras_prohibidas or "SEGUIMIENTO" in unidad:
        return False

    # --- LÓGICA DE ACEPTACIÓN ---
    # 1. Patrón para Romanos: Debe empezar con romano (I, V, X)
    # y NO ser solo la letra 'I' si es parte de una palabra larga (como 'INGENIERÍA')
    es_romano_valido = bool(re.match(r"^[IVXLCDM]+([\s\.,yY]|$)", unidad))

    # 2. Patrón para Dígitos: Empieza con número (1, 2, 3...)
    es_digito_valido = bool(re.match(r"^\d+", unidad))

    # 3. Caso especial: "UNIDAD 1", "UNIDAD I", etc.
    # Pero evitamos que sea solo la palabra "UNIDADES" (que ya filtramos arriba)
    es_palabra_unidad = unidad.startswith("UNIDAD") and "UNIDADES" not in unidad

    return es_romano_valido or es_digito_valido or es_palabra_unidad


def extraer_filas_seguimiento(table) -> list:
    filas = []
    for row in table.rows:
        celdas = [limpiar(cell.text) for cell in row.cells]
        if not fila_es_dato(celdas):
            continue
        # Columnas: 0-Periodo 1-Unidad 2-Captura 3-Aprobados 4-Reprobados
        #           5-Evidencias 6-%Avance 7-Producto No Conforme
        filas.append({
            "Periodo":                   celdas[0],
            "Unidades":                  celdas[1],
            "Captura de Calificaciones": celdas[2],
            "Aprobados":                 celdas[3],
            "Reprobados":                celdas[4],
            "Evidencias de Aprendizaje": celdas[5],
            "% De Avance":               celdas[6],
            "Producto No Conforme":      celdas[7],
        })
    return filas


# ---------------------------------------------------------------------------
# Extracción completa de un .docx
# ---------------------------------------------------------------------------

def extraer_docx(ruta: str) -> list:
    """Retorna lista de dicts, uno por unidad con datos encontrada."""
    doc = Document(ruta)
    encabezado = extraer_encabezado(doc)
    nombre_archivo = Path(ruta).name
    nombre_archivo = nombre_archivo.replace("sgi_temp_", "")

    registros = []
    for table in doc.tables:
        if not es_tabla_seguimiento(table):
            continue
        for fila in extraer_filas_seguimiento(table):
            registros.append({
                "Archivo": nombre_archivo,
                **encabezado,
                **fila,
            })

    # Si no hubo filas con datos, guardamos el encabezado igual
    if not registros:
        registros.append({
            "Archivo": nombre_archivo,
            **encabezado,
            "Periodo": "", "Unidades": "",
            "Captura de Calificaciones": "", "Aprobados": "",
            "Reprobados": "", "Evidencias de Aprendizaje": "",
            "% De Avance": "", "Producto No Conforme": "",
        })

    return registros


# ---------------------------------------------------------------------------
# Lectura de archivos: carpeta, ZIP, DOCX suelto
# ---------------------------------------------------------------------------

def procesar_docx(ruta: str, data_list: list):
    """Extrae datos de un .docx y los agrega a data_list."""
    print(f"  Procesando: {Path(ruta).name}")
    try:
        registros = extraer_docx(ruta)
        for r in registros:
            data_list.append([r.get(col, "") for col in ENCABEZADO_EXCEL])
        print(f"    → {len(registros)} fila(s) extraída(s)")
    except Exception as e:
        print(f"    [ERROR] {e}")


def convertir_doc_a_docx(doc_path: str) -> str:
    """Convierte un .doc legacy a .docx usando Word COM. Retorna la nueva ruta."""
    print(f"    Convirtiendo .doc a .docx: {Path(doc_path).name}")
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        nuevo_path = doc_path + "x"  # .doc → .docx
        doc = word.Documents.Open(os.path.abspath(doc_path))
        doc.SaveAs2(os.path.abspath(nuevo_path), FileFormat=16)  # 16 = docx
        doc.Close()
        word.Quit()
        return nuevo_path
    except Exception as e:
        print(f"    [ERROR] No se pudo convertir: {e}")
        return ""


def procesar_zip(zip_path: str, data_list: list):
    """Extrae todos los .docx/.doc dentro de un ZIP."""
    print(f"\nAbriendo ZIP: {Path(zip_path).name}")
    temp_dir = tempfile.gettempdir()

    with zipfile.ZipFile(zip_path, "r") as zf:
        for info in zf.infolist():
            nombre = info.filename
            basename = os.path.basename(nombre)

            if not (nombre.endswith(".docx") or nombre.endswith(".doc")):
                continue
            if basename.startswith("~$"):
                continue  # archivo temporal de Word

            temp_path = os.path.join(temp_dir, f"sgi_temp_{basename}")

            # Eliminar temporal previo si existe
            if os.path.exists(temp_path):
                try:
                    os.remove(temp_path)
                except PermissionError:
                    print(f"  [AVISO] No se pudo eliminar temporal previo: {temp_path}, omitiendo...")
                    continue

            with zf.open(nombre) as src, open(temp_path, "wb") as dst:
                dst.write(src.read())

            # Si es .doc (legacy), convertir primero con Word COM
            if nombre.endswith(".doc"):
                temp_path = convertir_doc_a_docx(temp_path)
                if not temp_path:
                    continue

            procesar_docx(temp_path, data_list)

            # Limpiar temporal
            if os.path.exists(temp_path):
                try:
                    os.remove(temp_path)
                except PermissionError:
                    print(f"  [AVISO] No se pudo eliminar: {temp_path}")


def procesar_ruta(ruta: str, data_list: list):
    """Despacha según si la ruta es carpeta, ZIP o DOCX."""
    p = Path(ruta)

    if p.is_dir():
        print(f"\nEscaneando carpeta: {ruta}")
        zips  = sorted(p.glob("**/*.zip"))
        docxs = sorted(p.glob("**/*.docx"))

        for z in zips:
            procesar_zip(str(z), data_list)

        for d in docxs:
            if not d.name.startswith("~$"):
                procesar_docx(str(d), data_list)

    elif p.suffix.lower() == ".zip":
        procesar_zip(str(p), data_list)

    elif p.suffix.lower() in (".docx", ".doc"):
        if p.suffix.lower() == ".doc":
            convertido = convertir_doc_a_docx(str(p))
            if convertido:
                procesar_docx(convertido, data_list)
        else:
            procesar_docx(str(p), data_list)

    else:
        print(f"[AVISO] Ruta no reconocida o no soportada: {ruta}")


# ---------------------------------------------------------------------------
# Generación del Excel
# ---------------------------------------------------------------------------

def crear_excel(data_list: list):
    """Crea datos.xlsx en el directorio actual con formato visual."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "SGI Datos"

    # Estilo encabezado
    fill_header = PatternFill("solid", fgColor="1F4E79")
    font_header = Font(bold=True, color="FFFFFF", size=11)
    alin_center = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for col_idx, campo in enumerate(data_list[0], start=1):
        cell = ws.cell(row=1, column=col_idx, value=campo)
        cell.fill = fill_header
        cell.font = font_header
        cell.alignment = alin_center

    # Datos con filas alternadas
    fill_par   = PatternFill("solid", fgColor="D9E1F2")
    fill_impar = PatternFill("solid", fgColor="FFFFFF")

    for row_idx, fila in enumerate(data_list[1:], start=2):
        fill = fill_par if row_idx % 2 == 0 else fill_impar
        for col_idx, valor in enumerate(fila, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=valor)
            cell.fill = fill
            cell.alignment = Alignment(vertical="center")

    # Ajustar ancho de columnas automáticamente
    for col in ws.columns:
        max_len = max((len(str(c.value)) if c.value else 0) for c in col)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 40)

    ws.row_dimensions[1].height = 40
    ws.freeze_panes = "A2"  # Encabezado fijo al hacer scroll

    ruta_excel = os.path.join(os.getcwd(), "datos.xlsx")
    wb.save(ruta_excel)
    print(f"\n✅ Excel generado: {ruta_excel}")
    print(f"   {len(data_list) - 1} fila(s) de datos escritas.")


# ---------------------------------------------------------------------------
# Punto de entrada
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    ruta_entrada = sys.argv[1] if len(sys.argv) > 1 else RUTA_POR_DEFECTO

    data_list = [ENCABEZADO_EXCEL]

    procesar_ruta(ruta_entrada, data_list)

    if len(data_list) > 1:
        crear_excel(data_list)
    else:
        print("\n[AVISO] No se encontraron datos para exportar.")