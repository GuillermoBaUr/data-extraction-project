"""
extract_sgi.py
--------------
Extracts fields from the SGI "Course Management Tracking" format
from individual .docx files, within .zip files, or entire folders.
Generates an Excel file with all data upon completion.

Usage:
    python extract_sgi.py                  <- uses '../data' folder by default
    python extract_sgi.py C:/folder/path   <- custom folder
    python extract_sgi.py C:/file.zip      <- single ZIP
    python extract_sgi.py C:/file.docx     <- single DOCX
"""

import os
import re
import sys
import zipfile
import tempfile
from pathlib import Path

import win32com.client
from docx import Document
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

DEFAULT_PATH = r"..\data"

EXCEL_HEADER = [
    "File", "Revision", "Teacher", "Subject",
    "Number of Enrolled Students",
    "Period", "Units", "Grade Entry",
    "Passed", "Failed", "Learning Evidence",
    "% Progress", "Non-Conforming Product",
]


# ---------------------------------------------------------------------------
# Text Helpers
# ---------------------------------------------------------------------------

def clean(text: str) -> str:
    """Removes filler underscores, line breaks, and multiple spaces."""
    text = text.replace("\n", " ").replace("\r", " ").strip()
    text = re.sub(r"_{2,}", "", text)
    text = re.sub(r"\s{2,}", " ", text)
    return text.strip()


def extract_field(text: str, pattern: str) -> str:
    """Searches for 'pattern: VALUE' tolerating case variations and spaces."""
    match = re.search(pattern + r"[:\s]+(.+?)(?=\s{2,}|$)", text, re.IGNORECASE)
    return clean(match.group(1)) if match else ""


def is_placeholder(value: str) -> bool:
    """Returns True if the value is an unfilled instructional marker: (2), (3), etc."""
    return bool(re.fullmatch(r"\(\d{1,2}\)", value.strip()))


# ---------------------------------------------------------------------------
# Header Extraction (Revision, Teacher, Subject, Students)
# ---------------------------------------------------------------------------

def extract_header(doc: Document) -> dict:
    data = {
        "Revision":                     "",
        "Teacher":                      "",
        "Subject":                      "",
        "Number of Enrolled Students":  "",
    }

    # Revision is in the header table of the document section
    for section in doc.sections:
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    match = re.search(r"Revisi[oó]n[:\s]*(\d+)", cell.text, re.IGNORECASE)
                    if match:
                        data["Revision"] = match.group(1)

    # Teacher, Subject, and Students in free text paragraphs
    waiting_for_teacher_name = False
    previous_text = ""  # Stores the previous paragraph

    for paragraph in doc.paragraphs:
        text = paragraph.text
        if not text.strip():
            previous_text = ""
            continue

        if not data["Teacher"]:
            combined_text = previous_text + text

            # Case 1: TEACHER (DOCENTE) complete in current paragraph
            if re.search(r"DOCENTE\s*:?", text, re.IGNORECASE):
                m = re.search(r"DOCENTE\s*:?\s*_*\s*(?P<name>[^_\n\r\t]+?)(?=_+|\s*$)", text, re.IGNORECASE)
                if m and m.group("name").strip():
                    val = re.sub(r"\s{2,}", " ", m.group("name").strip())
                    if not is_placeholder(val):
                        data["Teacher"] = val.upper()
                else:
                    waiting_for_teacher_name = True

            # Case 2: TEACHER split between paragraphs
            elif re.search(r"DOCENTE\s*:?:", combined_text, re.IGNORECASE):
                m2 = re.search(r"DOCENTE\s*:?\s*_*\s*(?P<name>[^_\n\r\t]+?)(?=_+|\s*$)", combined_text, re.IGNORECASE)
                if m2 and m2.group("name").strip():
                    val = re.sub(r"\s{2,}", " ", m2.group("name").strip())
                    if not is_placeholder(val):
                        data["Teacher"] = val.upper()

            # Case 3: Name in the next paragraph (flag from previous loop)
            elif waiting_for_teacher_name and not re.search(r"ASIGNATURA|DEPARTAMENTO|CANTIDAD", text, re.IGNORECASE):
                val = re.sub(r"\s{2,}", " ", text.strip().strip("_").strip())
                if not is_placeholder(val):
                    data["Teacher"] = val
                waiting_for_teacher_name = False

        if not data["Subject"]:
            m = re.search(r"ASIGNATURA(?:\s+Y\s+GRUPO)?\s*:\s*(?P<name>.+?)(?=CANTIDAD|\s{5,}|$)", text, re.IGNORECASE)
            if m:
                val = m.group("name").replace("_", " ").strip()
                val = re.sub(r"\s{2,}", " ", val)
                if not is_placeholder(val):
                    data["Subject"] = val.upper()

        if not data["Number of Enrolled Students"]:
            m = re.search(r"CANTIDAD\s+DE\s+ESTUDIANTES?\s+INSCRITOS?[:\s]+(?:\(\d+\))?[\s_]*(?P<num>\d+)", text, re.IGNORECASE)
            if m:
                val = m.group("num")
                if not is_placeholder(val):
                    data["Number of Enrolled Students"] = val

        previous_text = text

    return data

# ---------------------------------------------------------------------------
# Tracking Table Extraction
# ---------------------------------------------------------------------------

_TABLE_KEYWORDS = {"UNIDADES", "APROBADOS", "REPROBADOS", "AVANCE", "CONFORME"}


def is_tracking_table(table) -> bool:
    """Identifies the data table by keywords in its headers."""
    header_text = " ".join(cell.text.upper() for cell in table.rows[0].cells)
    return sum(1 for kw in _TABLE_KEYWORDS if kw in header_text) >= 3


def row_is_data(cells: list) -> bool:
    """
    Accepts actual data rows and excludes table header rows.
    """
    if not cells or len(cells) < 2:
        return False

    unit_cell = cells[1].strip().upper()

    # --- EXCLUSION FILTER ---
    forbidden_words = {"UNIDADES", "PERIODO", "APROBADOS", "REPROBADOS", "CAPACITACIÓN", "UNITS", "PERIOD"}
    if unit_cell in forbidden_words or "SEGUIMIENTO" in unit_cell or "TRACKING" in unit_cell:
        return False

    # --- ACCEPTANCE LOGIC ---
    # 1. Roman Numerals Pattern (I, V, X...)
    is_valid_roman = bool(re.match(r"^[IVXLCDM]+([\s\.,yY]|$)", unit_cell))

    # 2. Digits Pattern (1, 2, 3...)
    is_valid_digit = bool(re.match(r"^\d+", unit_cell))

    # 3. Special case: "UNIT 1", "UNIDAD I", etc.
    is_unit_word = (unit_cell.startswith("UNIDAD") or unit_cell.startswith("UNIT")) and "UNIDADES" not in unit_cell

    is_topic = bool(re.match(r"^(TEMA|TOPIC)\s+[IVXLCDM\d]+", unit_cell))

    return is_valid_roman or is_valid_digit or is_unit_word or is_topic


def extract_tracking_rows(table) -> list:
    rows = []
    for row in table.rows:
        cells = [clean(cell.text) for cell in row.cells]
        if not row_is_data(cells):
            continue
        # Columns: 0-Period 1-Unit 2-Entry 3-Passed 4-Failed 5-Evidence 6-%Progress 7-Non-conforming
        rows.append({
            "Period":                   cells[0],
            "Units":                    cells[1],
            "Grade Entry":              cells[2],
            "Passed":                   cells[3],
            "Failed":                   cells[4],
            "Learning Evidence":        cells[5],
            "% Progress":               cells[6],
            "Non-Conforming Product":   cells[7],
        })
    return rows


# ---------------------------------------------------------------------------
# Complete .docx Extraction
# ---------------------------------------------------------------------------

def extract_docx(path: str) -> list:
    """Returns a list of dicts, one per found data unit."""
    doc = Document(path)
    header_data = extract_header(doc)
    file_name = Path(path).name
    file_name = file_name.replace("sgi_temp_", "")

    records = []
    for table in doc.tables:
        if not is_tracking_table(table):
            continue
        for row in extract_tracking_rows(table):
            records.append({
                "File": file_name,
                **header_data,
                **row,
            })

    # If no data rows were found, save the header anyway
    if not records:
        records.append({
            "File": file_name,
            **header_data,
            "Period": "", "Units": "",
            "Grade Entry": "", "Passed": "",
            "Failed": "", "Learning Evidence": "",
            "% Progress": "", "Non-Conforming Product": "",
        })

    return records


# ---------------------------------------------------------------------------
# File Reading: Folder, ZIP, Loose DOCX
# ---------------------------------------------------------------------------

def process_docx(path: str, data_list: list):
    """Extracts data from a .docx and adds it to data_list."""
    print(f"  Processing: {Path(path).name}")
    try:
        records = extract_docx(path)
        for r in records:
            # Map translated dict keys to the English EXCEL_HEADER order
            row_to_append = []
            mapping = {
                "File": "File", "Revision": "Revision", "Teacher": "Teacher", "Subject": "Subject",
                "Number of Enrolled Students": "Number of Enrolled Students",
                "Period": "Period", "Units": "Units", "Grade Entry": "Grade Entry",
                "Passed": "Passed", "Failed": "Failed", "Learning Evidence": "Learning Evidence",
                "% Progress": "% Progress", "Non-Conforming Product": "Non-Conforming Product"
            }
            # This follows the order in EXCEL_HEADER
            data_list.append([r.get(mapping.get(col, col), "") for col in EXCEL_HEADER])
        print(f"    → {len(records)} row(s) extracted")
    except Exception as e:
        print(f"    [ERROR] {e}")


def convert_doc_to_docx(doc_path: str) -> str:
    """Converts a legacy .doc to .docx using Word COM. Returns the new path."""
    print(f"    Converting .doc to .docx: {Path(doc_path).name}")
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        new_path = doc_path + "x"  # .doc → .docx
        doc = word.Documents.Open(os.path.abspath(doc_path))
        doc.SaveAs2(os.path.abspath(new_path), FileFormat=16)  # 16 = docx
        doc.Close()
        word.Quit()
        return new_path
    except Exception as e:
        print(f"    [ERROR] Could not convert: {e}")
        return ""


def process_zip(zip_path: str, data_list: list):
    """Extracts all .docx/.doc inside a ZIP."""
    print(f"\nOpening ZIP: {Path(zip_path).name}")
    temp_dir = tempfile.gettempdir()

    with zipfile.ZipFile(zip_path, "r") as zf:
        for info in zf.infolist():
            name = info.filename
            basename = os.path.basename(name)

            if not (name.endswith(".docx") or name.endswith(".doc")):
                continue
            if basename.startswith("~$"):
                continue  # Word temp file

            temp_path = os.path.join(temp_dir, f"sgi_temp_{basename}")

            # Delete previous temp if exists
            if os.path.exists(temp_path):
                try:
                    os.remove(temp_path)
                except PermissionError:
                    print(f"  [NOTICE] Could not delete previous temp: {temp_path}, skipping...")
                    continue

            with zf.open(name) as src, open(temp_path, "wb") as dst:
                dst.write(src.read())

            # If legacy .doc, convert first
            if name.endswith(".doc"):
                temp_path = convert_doc_to_docx(temp_path)
                if not temp_path:
                    continue

            process_docx(temp_path, data_list)

            # Clean up temp
            if os.path.exists(temp_path):
                try:
                    os.remove(temp_path)
                except PermissionError:
                    print(f"  [NOTICE] Could not delete: {temp_path}")


def process_path(path: str, data_list: list):
    """Dispatches based on whether the path is a folder, ZIP, or loose DOCX."""
    p = Path(path)

    if p.is_dir():
        print(f"\nScanning folder: {path}")
        zips  = sorted(p.glob("**/*.zip"))
        docxs = sorted(p.glob("**/*.docx"))

        for z in zips:
            process_zip(str(z), data_list)

        for d in docxs:
            if not d.name.startswith("~$"):
                process_docx(str(d), data_list)

    elif p.suffix.lower() == ".zip":
        process_zip(str(p), data_list)

    elif p.suffix.lower() in (".docx", ".doc"):
        if p.suffix.lower() == ".doc":
            converted = convert_doc_to_docx(str(p))
            if converted:
                process_docx(converted, data_list)
        else:
            process_docx(str(p), data_list)

    else:
        print(f"[NOTICE] Path not recognized or supported: {path}")


# ---------------------------------------------------------------------------
# Excel Generation
# ---------------------------------------------------------------------------

def create_excel(data_list: list):
    """Creates data.xlsx in the current directory with visual formatting."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "SGI Data"

    # Header style
    header_fill = PatternFill("solid", fgColor="1F4E79")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for col_idx, field in enumerate(data_list[0], start=1):
        cell = ws.cell(row=1, column=col_idx, value=field)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center_align

    # Data with alternating rows
    even_fill = PatternFill("solid", fgColor="D9E1F2")
    odd_fill  = PatternFill("solid", fgColor="FFFFFF")

    for row_idx, row_data in enumerate(data_list[1:], start=2):
        fill = even_fill if row_idx % 2 == 0 else odd_fill
        for col_idx, value in enumerate(row_data, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.fill = fill
            cell.alignment = Alignment(vertical="center")

    # Auto-adjust column width
    for col in ws.columns:
        max_len = max((len(str(c.value)) if c.value else 0) for c in col)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 40)

    ws.row_dimensions[1].height = 40
    ws.freeze_panes = "A2"  # Freeze header on scroll

    excel_path = os.path.join(os.getcwd(), "data.xlsx")
    wb.save(excel_path)
    print(f"\n✅ Excel generated: {excel_path}")
    print(f"   {len(data_list) - 1} data row(s) written.")


# ---------------------------------------------------------------------------
# Main Entry Point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    input_path = sys.argv[1] if len(sys.argv) > 1 else DEFAULT_PATH

    data_list = [EXCEL_HEADER]

    process_path(input_path, data_list)

    if len(data_list) > 1:
        create_excel(data_list)
    else:
        print("\n[NOTICE] No data found to export.")